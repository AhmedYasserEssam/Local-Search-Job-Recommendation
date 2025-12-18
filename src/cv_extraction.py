import re
import fitz
import ollama
from docx import Document
from dataclasses import dataclass
from typing import List
from datetime import datetime

@dataclass
class CVData:
    raw_text: str
    skills: List[str]
    experience_years: float


def extract_text(file_path: str) -> str:
    def from_pdf():
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
        return text.strip()

    def from_docx():
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)

    if file_path.lower().endswith('.pdf'):
        return from_pdf()
    elif file_path.lower().endswith('.docx'):
        return from_docx()


import re
from datetime import datetime
from typing import List, Tuple
import ollama


def extract_experience_years(full_text: str, model: str = "llama3.2") -> float:

    explicit_pattern = r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp|professional)'
    match = re.search(explicit_pattern, full_text.lower())
    if match:
        years = float(match.group(1))
        if 0 < years < 60:
            return round(years, 1)

    prompt = f"""
Extract ALL job date ranges from the Work Experience section.

Return ONLY date ranges, one per line, in this exact format:
Month YYYY to Month YYYY

Examples:
Oct 2022 to Current
Sept 2021 to Aug 2022

Rules:
- Always include BOTH start and end dates with FULL format (Month YYYY)
- Use "Current" or "Present" if still working
- One date range per line

CV:
{full_text[:6000]}

Dates:
"""

    response = ollama.generate(
        model=model,
        prompt=prompt,
        options={"temperature": 0, "num_predict": 300}
    )

    text = response.get("response", "").lower()
    text = text.replace('�', '-').replace('�', '-').replace('–', '-').replace('—', '-').replace('now', 'current')

    current = datetime.now()
    current_month, current_year = current.month, current.year

    month_map = {
        'jan': 1, 'january': 1, 'feb': 2, 'february': 2,
        'mar': 3, 'march': 3, 'apr': 4, 'april': 4,
        'may': 5, 'jun': 6, 'june': 6,
        'jul': 7, 'july': 7, 'aug': 8, 'august': 8,
        'sep': 9, 'sept': 9, 'september': 9,
        'oct': 10, 'october': 10, 'nov': 11, 'november': 11,
        'dec': 12, 'december': 12
    }

    ranges: List[Tuple[int, int]] = []

    numeric_pattern = r'(\d{1,2})/(\d{4})\s*(?:to|-)\s*(?:(\d{1,2})/(\d{4})|(present|current|now))'

    month_pattern = (
        r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
        r'jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|'
        r'nov(?:ember)?|dec(?:ember)?)\s+(\d{4})\s*(?:to|-)\s*'
        r'(?:(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|'
        r'jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|'
        r'oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+(\d{4})|(present|current|now))'
    )

    year_pattern = r'(\d{4})\s*(?:to|-)\s*(\d{4}|present|current|now)'

    for m in re.findall(numeric_pattern, text):
        sm, sy = int(m[0]), int(m[1])
        if m[4]:
            em, ey = current_month, current_year
        else:
            em, ey = int(m[2]), int(m[3])

        start = sy * 12 + sm
        end = ey * 12 + em
        if 0 < end - start < 720:
            ranges.append((start, end))

    for m in re.findall(month_pattern, text):
        sm = month_map[m[0][:3]]
        sy = int(m[1])
        if m[4]:
            em, ey = current_month, current_year
        else:
            em = month_map[m[2][:3]]
            ey = int(m[3])

        start = sy * 12 + sm
        end = ey * 12 + em
        if 0 < end - start < 720:
            ranges.append((start, end))

    for m in re.findall(year_pattern, text):
        sy = int(m[0])
        if m[1] in ("present", "current"):
            ey = current_year
        else:
            ey = int(m[1])

        start = sy * 12 + 1
        end = ey * 12 + 12
        if 0 < end - start < 720:
            ranges.append((start, end))
    if not ranges:
        return 0.0

    ranges.sort()
    merged = [ranges[0]]

    for start, end in ranges[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))

    total_months = sum(end - start for start, end in merged)

    return round(total_months / 12, 1)

def extract_skills(full_text: str, model: str = "llama3.2") -> List[str]:
    prompt = f"""Extract ALL skills from this CV/resume.

Return ONLY a comma-separated list of skills, nothing else.

Rules:
- Include ALL skills: technical, soft skills, tools, technologies, methodologies
- Include domain-specific skills from ANY field
- Extract from everywhere: skills section, job descriptions, projects, education, certifications
- Each skill should be concise (1-4 words)
- Do not include job titles, company names, or dates
- Do not add explanations or categories

CV Text:
{full_text[:8000]}

Skills (comma-separated):"""

    response = ollama.generate(
        model=model,
        prompt=prompt,
        options={"temperature": 0.2, "num_predict": 1000}
    )
    
    response_text = response.get("response", "").strip()
    
    if not response_text:
        return []
    
    skills = []
    for skill in response_text.split(','):
        cleaned = skill.strip().strip('-').strip('•').strip('"').strip("'")
        if cleaned and 1 < len(cleaned) < 60 and not cleaned.lower().startswith('skill'):
            skills.append(cleaned)
    
    seen = set()
    unique_skills = []
    for skill in skills:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            seen.add(skill_lower)
            unique_skills.append(skill)
    
    return unique_skills


def extract_cv_data(file_path: str, model: str = "llama3.2") -> CVData:
    raw_text = extract_text(file_path)
    skills = extract_skills(raw_text, model)
    experience_years = extract_experience_years(raw_text, model)

    return CVData(
        raw_text=raw_text,
        skills=skills,
        experience_years=experience_years,
    )
